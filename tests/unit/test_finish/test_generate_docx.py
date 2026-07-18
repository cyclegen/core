"""test_generate_docx.py — generate_docx テンプレート別出力検証"""

from pathlib import Path

import pytest
from docx import Document

from cyclegen.finish.converter import generate_docx, parse_markdown
from cyclegen.finish.templates import list_templates, load_template

SAMPLE_MD = """\
# Sample Document

## Section One

Body text with **bold** and *italic*.

- List item 1
- List item 2

| Header A | Header B |
|----------|----------|
| Cell 1   | Cell 2   |

> A blockquote

```python
print("hello")
```

---

1. Ordered item 1
2. Ordered item 2
"""

ALL_TEMPLATE_NAMES = [t["name"] for t in list_templates()]


@pytest.fixture
def elements():
    return parse_markdown(SAMPLE_MD)


class TestGenerateDocxPerTemplate:
    @pytest.mark.parametrize("template_name", ALL_TEMPLATE_NAMES)
    def test_generates_valid_docx(self, elements, template_name, tmp_path):
        tmpl = load_template(template_name)
        out = tmp_path / f"output_{template_name}.docx"
        result = generate_docx(elements, tmpl, out)

        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

        # python-docx で読み返して壊れていないことを確認
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0


class TestGenerateDocxContent:
    def test_heading_present(self, elements, tmp_path):
        tmpl = load_template("executive")
        out = tmp_path / "test.docx"
        generate_docx(elements, tmpl, out)

        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "Sample Document" in texts
        assert "Section One" in texts

    def test_table_present(self, elements, tmp_path):
        tmpl = load_template("executive")
        out = tmp_path / "test.docx"
        generate_docx(elements, tmpl, out)

        doc = Document(str(out))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Header A"

    def test_output_path_as_string(self, elements, tmp_path):
        tmpl = load_template("minimal")
        out = str(tmp_path / "string_path.docx")
        result = generate_docx(elements, tmpl, out)
        assert Path(result).exists()

    def test_empty_elements(self, tmp_path):
        tmpl = load_template("executive")
        out = tmp_path / "empty.docx"
        result = generate_docx([], tmpl, out)
        assert result.exists()

    def test_image_missing_placeholder(self, tmp_path):
        """存在しない画像はプレースホルダーになる"""
        elements = parse_markdown("![alt text](nonexistent.png)")
        tmpl = load_template("executive")
        out = tmp_path / "img_test.docx"
        generate_docx(elements, tmpl, out)

        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("画像" in t for t in texts)

    def test_toc_generated(self, tmp_path):
        md = "<!-- toc -->\n\n## Chapter 1\n\n### Section 1.1\n\n## Chapter 2"
        elements = parse_markdown(md)
        tmpl = load_template("executive")
        out = tmp_path / "toc_test.docx"
        generate_docx(elements, tmpl, out)

        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert "目次" in texts

    def test_pagebreak(self, tmp_path):
        md = "## Before\n\n<!-- pagebreak -->\n\n## After"
        elements = parse_markdown(md)
        tmpl = load_template("executive")
        out = tmp_path / "pb_test.docx"
        generate_docx(elements, tmpl, out)
        assert out.exists()
