"""Markdown → docx 変換コアロジック

元ソース: _CycleGen_Core/scripts/finish/cyclegen_finish.py
CLIの main() は移動対象外。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor


# ============================================================
# ユーティリティ
# ============================================================


def hex_to_rgb(hex_str: str) -> RGBColor:
    """'1B2A4A' -> RGBColor"""
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


# ============================================================
# Markdown パーサー
# ============================================================


def parse_markdown(content: str) -> list[dict]:
    """Markdownを構造化された要素のリストに変換する"""
    elements: list[dict] = []
    content = content.replace("\r\n", "\n").replace("\r", "\n")
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # ページ区切りマーカー
        if line.strip() == "<!-- pagebreak -->":
            elements.append({"type": "pagebreak"})
            i += 1
            continue

        # 目次マーカー
        if line.strip() == "<!-- toc -->":
            elements.append({"type": "toc"})
            i += 1
            continue

        # コードブロック
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            elements.append({
                "type": "code_block",
                "language": lang,
                "code": "\n".join(code_lines),
            })
            continue

        # 見出し
        m = re.match(r"^(#{1,5})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            elements.append({"type": "heading", "level": level, "text": m.group(2)})
            i += 1
            continue

        # 水平線
        if line.strip() in ("---", "***", "___"):
            elements.append({"type": "hr"})
            i += 1
            continue

        # 画像
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if img_match:
            elements.append({
                "type": "image",
                "alt": img_match.group(1),
                "path": img_match.group(2),
            })
            i += 1
            continue

        # テーブル
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            elements.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 引用ブロック
        if line.strip().startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            elements.append({"type": "blockquote", "text": "\n".join(quote_lines)})
            continue

        # リスト項目（ネスト対応）
        if re.match(r"^(\s*)- ", line):
            items = []
            while i < len(lines) and re.match(r"^(\s*)- ", lines[i]):
                indent_match = re.match(r"^(\s*)- (.+)$", lines[i])
                if indent_match:
                    indent_level = len(indent_match.group(1)) // 2
                    items.append({
                        "text": indent_match.group(2).strip(),
                        "level": indent_level,
                    })
                i += 1
            elements.append({"type": "list", "items": items})
            continue

        # 番号付きリスト
        if re.match(r"^\s*\d+\.\s", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s", lines[i]):
                indent_match = re.match(r"^(\s*)\d+\.\s*(.+)$", lines[i])
                if indent_match:
                    indent_level = len(indent_match.group(1)) // 2
                    items.append({
                        "text": indent_match.group(2).strip(),
                        "level": indent_level,
                    })
                i += 1
            elements.append({"type": "ordered_list", "items": items})
            continue

        # 通常の段落
        para_lines = []
        while i < len(lines) and lines[i].strip() \
                and not re.match(r"^#{1,5}\s", lines[i]) \
                and lines[i].strip() not in ("---", "***", "___") \
                and lines[i].strip() != "<!-- pagebreak -->" \
                and lines[i].strip() != "<!-- toc -->" \
                and not lines[i].strip().startswith("```") \
                and not lines[i].strip().startswith("> ") \
                and not re.match(r"^(\s*)- ", lines[i]) \
                and not re.match(r"^\s*\d+\.\s", lines[i]) \
                and not re.match(r"^!\[", lines[i].strip()) \
                and not ("|" in lines[i] and i + 1 < len(lines)
                         and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip())):
            raw_line = lines[i]
            clean_line = re.sub(r"<br\s*/?\s*>", "\n", raw_line.strip())
            if raw_line.rstrip("\n").endswith("  "):
                clean_line = clean_line.rstrip() + "\n"
            para_lines.append(clean_line)
            i += 1
        if para_lines:
            elements.append({"type": "paragraph", "text": "\n".join(
                pl.rstrip("\n") for pl in para_lines
            )})

    return elements


# ============================================================
# テキストレンダリング（Bold/Italic/Link対応）
# ============================================================


def add_hyperlink(paragraph, url: str, text: str, font_name: str, font_size: float):
    """段落にハイパーリンクを追加する"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
    )
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr>'
        f'<w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="2E86C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_formatted_text(paragraph, text: str, style: dict, doc=None):
    """**bold**、*italic*、[text](url) を含むテキストをparagraphに追加する。
    \\n は行内改行として処理する。"""
    font_name = style.get("font", "Arial")
    font_size = style.get("size", 10.5)
    font_color = style.get("color", "333333")

    line_segments = text.split("\n")
    for seg_idx, segment in enumerate(line_segments):
        if seg_idx > 0:
            run = paragraph.add_run()
            run.add_break()
        if not segment:
            continue

        parts = re.split(r"(\[([^\]]+)\]\(([^)]+)\)|\*\*[^*]+\*\*|\*[^*]+\*)", segment)

        idx = 0
        while idx < len(parts):
            part = parts[idx]
            if not part:
                idx += 1
                continue

            link_match = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                if doc:
                    add_hyperlink(paragraph, link_url, link_text, font_name, font_size)
                else:
                    run = paragraph.add_run(link_text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = hex_to_rgb("2E86C1")
                    run.underline = True
                idx += 3
                continue

            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)

            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = hex_to_rgb(font_color)
            idx += 1


# ============================================================
# docx 生成
# ============================================================


def set_cell_shading(cell, color_hex: str):
    """テーブルセルの背景色を設定する"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def collect_headings(elements: list[dict]) -> list[dict]:
    """目次用にH2-H5見出しを収集する"""
    return [
        {"level": elem["level"], "text": elem["text"]}
        for elem in elements
        if elem["type"] == "heading" and elem["level"] in (2, 3, 4, 5)
    ]


def _copy_section_layout(source_section, target_section):
    """セクションのページレイアウトを引き継ぐ"""
    target_section.page_width = source_section.page_width
    target_section.page_height = source_section.page_height
    target_section.top_margin = source_section.top_margin
    target_section.bottom_margin = source_section.bottom_margin
    target_section.left_margin = source_section.left_margin
    target_section.right_margin = source_section.right_margin


def generate_docx(
    elements: list[dict],
    template: dict,
    output_path: Path | str,
    base_dir: Path | None = None,
) -> Path:
    """構造化要素からdocxを生成する。

    Returns:
        生成されたdocxファイルのパス
    """
    output_path = Path(output_path)
    doc = Document()

    # Normalスタイルのデフォルトをテンプレートのbody設定で上書き
    body_cfg = template.get("body", {})
    normal_style = doc.styles["Normal"]
    normal_style.font.name = body_cfg.get("font", "Arial")
    normal_style.font.size = Pt(body_cfg.get("size", 10.5))
    normal_style.font.color.rgb = hex_to_rgb(body_cfg.get("color", "333333"))
    normal_style.paragraph_format.line_spacing = body_cfg.get("line_spacing", 1.0)
    normal_style.paragraph_format.space_before = Pt(body_cfg.get("space_before", 4))
    normal_style.paragraph_format.space_after = Pt(body_cfg.get("space_after", 4))

    # ページ設定
    page = template.get("page", {})
    for section in doc.sections:
        if page.get("width"):
            section.page_width = Mm(page["width"])
        if page.get("height"):
            section.page_height = Mm(page["height"])
        section.top_margin = Cm(page.get("margin_top", 2.5))
        section.bottom_margin = Cm(page.get("margin_bottom", 2.5))
        section.left_margin = Cm(page.get("margin_left", 2.5))
        section.right_margin = Cm(page.get("margin_right", 2.5))

    # フッター設定（ページ番号）
    footer_cfg = template.get("footer", {})
    if footer_cfg.get("enabled", True):
        font_name = footer_cfg.get("font", "Arial")
        font_size = footer_cfg.get("size", 8)
        font_color = footer_cfg.get("color", "999999")
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # PAGE フィールドでページ番号を挿入
            rPr_xml = (
                f'<w:rPr {nsdecls("w")}>'
                f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
                f'<w:sz w:val="{int(font_size * 2)}"/>'
                f'<w:color w:val="{font_color}"/>'
                f'</w:rPr>'
            )
            run_begin = parse_xml(
                f'<w:r {nsdecls("w")}>{rPr_xml}'
                f'<w:fldChar w:fldCharType="begin"/></w:r>'
            )
            run_instr = parse_xml(
                f'<w:r {nsdecls("w")}>{rPr_xml}'
                f'<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
            )
            run_separate = parse_xml(
                f'<w:r {nsdecls("w")}>{rPr_xml}'
                f'<w:fldChar w:fldCharType="separate"/></w:r>'
            )
            run_default = parse_xml(
                f'<w:r {nsdecls("w")}>{rPr_xml}'
                f'<w:t>1</w:t></w:r>'
            )
            run_end = parse_xml(
                f'<w:r {nsdecls("w")}>{rPr_xml}'
                f'<w:fldChar w:fldCharType="end"/></w:r>'
            )
            p._p.append(run_begin)
            p._p.append(run_instr)
            p._p.append(run_separate)
            p._p.append(run_default)
            p._p.append(run_end)

    # 目次用の見出しを事前収集
    toc_headings = collect_headings(elements)

    # 要素をレンダリング
    for elem in elements:
        _render_element(doc, elem, template, toc_headings, base_dir)

    doc.save(str(output_path))
    return output_path


def _render_element(doc, elem: dict, template: dict, toc_headings: list[dict], base_dir: Path | None):
    """単一要素をdocにレンダリングする"""
    etype = elem["type"]

    if etype == "heading":
        _render_heading(doc, elem, template)
    elif etype == "paragraph":
        _render_paragraph(doc, elem, template)
    elif etype == "list":
        _render_list(doc, elem, template, ordered=False)
    elif etype == "ordered_list":
        _render_list(doc, elem, template, ordered=True)
    elif etype == "table":
        _render_table(doc, elem, template)
    elif etype == "hr":
        _render_hr(doc, template)
    elif etype == "code_block":
        _render_code_block(doc, elem, template)
    elif etype == "blockquote":
        _render_blockquote(doc, elem, template)
    elif etype == "image":
        _render_image(doc, elem, base_dir)
    elif etype == "pagebreak":
        doc.add_page_break()
    elif etype == "toc":
        _render_toc(doc, template, toc_headings)


def _render_heading(doc, elem: dict, template: dict):
    level = elem["level"]
    odd_page_start = template.get("page", {}).get("odd_page_start", False)

    if level == 1:
        if odd_page_start:
            new_section = doc.add_section(WD_SECTION_START.ODD_PAGE)
            prev = doc.sections[-2] if len(doc.sections) > 1 else doc.sections[0]
            _copy_section_layout(prev, new_section)
        cfg = template.get("title", {})
        p = doc.add_paragraph()
        run = p.add_run(elem["text"])
        run.bold = cfg.get("bold", True)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 24))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "222222"))
        alignment = cfg.get("alignment", "left")
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif level == 2:
        if odd_page_start:
            new_section = doc.add_section(WD_SECTION_START.ODD_PAGE)
            prev = doc.sections[-2] if len(doc.sections) > 1 else doc.sections[0]
            _copy_section_layout(prev, new_section)
        cfg = template.get("heading2", {})
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(cfg.get("space_before", 18))
        p.paragraph_format.space_after = Pt(cfg.get("space_after", 6))
        run = p.add_run(elem["text"])
        run.bold = cfg.get("bold", True)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 14))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "1B2A4A"))

        if cfg.get("border_bottom"):
            border_color = cfg.get("border_color", "C49A6C")
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="4" '
                f'w:color="{border_color}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    elif level == 3:
        cfg = template.get("heading3", {})
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(cfg.get("space_before", 12))
        p.paragraph_format.space_after = Pt(cfg.get("space_after", 4))
        run = p.add_run(elem["text"])
        run.bold = cfg.get("bold", True)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 12))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "555555"))

    elif level == 4:
        cfg = template.get("heading4", {})
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(cfg.get("space_before", 10))
        p.paragraph_format.space_after = Pt(cfg.get("space_after", 3))
        run = p.add_run(elem["text"])
        run.bold = cfg.get("bold", True)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 11))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "555555"))

    elif level == 5:
        cfg = template.get("heading5", {})
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(cfg.get("space_before", 8))
        p.paragraph_format.space_after = Pt(cfg.get("space_after", 2))
        run = p.add_run(elem["text"])
        run.bold = cfg.get("bold", False)
        run.italic = cfg.get("italic", True)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 10.5))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "666666"))


def _render_paragraph(doc, elem: dict, template: dict):
    cfg = template.get("body", {})
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = cfg.get("line_spacing", 1.0)
    p.paragraph_format.space_before = Pt(cfg.get("space_before", 4))
    p.paragraph_format.space_after = Pt(cfg.get("space_after", 4))
    add_formatted_text(p, elem["text"], cfg, doc)


def _render_list(doc, elem: dict, template: dict, *, ordered: bool):
    cfg = template.get("list", template.get("body", {}))
    counter = 1
    for item in elem["items"]:
        if isinstance(item, dict):
            text = item["text"]
            level = item.get("level", 0)
        else:
            text = item
            level = 0
        p = doc.add_paragraph()
        base_indent = cfg.get("indent", 1.0)
        p.paragraph_format.left_indent = Cm(base_indent + level * 0.8)

        if ordered:
            prefix = f"{counter}. "
            counter += 1
        else:
            prefix = "◦ " if level > 0 else "• "

        run = p.add_run(prefix)
        run.font.name = cfg.get("font", "Arial")
        run.font.size = Pt(cfg.get("size", 10.5))
        run.font.color.rgb = hex_to_rgb(cfg.get("color", "333333"))
        add_formatted_text(p, text, cfg, doc)


def _render_table(doc, elem: dict, template: dict):
    table_cfg = template.get("table", {})
    headers = elem["headers"]
    rows = elem["rows"]
    cols = max(len(headers), max((len(r) for r in rows), default=0))

    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_color = table_cfg.get("border_color", "CCCCCC")
    border_size = table_cfg.get("border_size", 4)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'<w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'<w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    header_style = {
        "font": table_cfg.get("header_font", "Arial"),
        "size": table_cfg.get("header_size", 10),
        "color": table_cfg.get("header_text", "FFFFFF"),
    }
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, table_cfg.get("header_bg", "1B2A4A"))
        p = cell.paragraphs[0]
        add_formatted_text(p, header_text, header_style, doc)
        for run in p.runs:
            run.bold = True

    body_style = {
        "font": table_cfg.get("body_font", "Arial"),
        "size": table_cfg.get("body_size", 10),
        "color": table_cfg.get("body_color", "333333"),
    }
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < cols:
                cell = table.rows[i + 1].cells[j]
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text, body_style, doc)


def _render_hr(doc, template: dict):
    accent = template.get("colors", {}).get("accent", "CCCCCC")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" '
        f'w:color="{accent}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _render_code_block(doc, elem: dict, template: dict):
    code_cfg = template.get("code", {})
    bg_color = code_cfg.get("bg", "F5F5F5")
    font_name = code_cfg.get("font", "Courier New")
    font_size = code_cfg.get("size", 9)
    font_color = code_cfg.get("color", "333333")

    for code_line in elem["code"].split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.0

        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
        pPr.append(shading)

        run = p.add_run(code_line if code_line else " ")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = hex_to_rgb(font_color)


def _render_blockquote(doc, elem: dict, template: dict):
    quote_cfg = template.get("blockquote", template.get("body", {}))
    accent = template.get("colors", {}).get("accent", "CCCCCC")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = quote_cfg.get("line_spacing", 1.5)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" '
        f'w:color="{accent}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    for qi, quote_line in enumerate(elem["text"].split("\n")):
        if qi > 0:
            br_run = p.add_run()
            br_run.add_break()
        if quote_line:
            runs_before = len(p.runs)
            add_formatted_text(p, quote_line, quote_cfg, doc)
            for run in p.runs[runs_before:]:
                run.italic = True


def _render_image(doc, elem: dict, base_dir: Path | None):
    img_path = elem["path"]
    if base_dir and not Path(img_path).is_absolute():
        img_path = str(base_dir / img_path)

    if Path(img_path).exists():
        current_section = doc.sections[-1]
        available_width = current_section.page_width - current_section.left_margin - current_section.right_margin
        doc.add_picture(img_path, width=available_width)
        last_p = doc.paragraphs[-1]
        last_p.paragraph_format.space_after = Pt(0)
        if elem.get("alt"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(elem["alt"])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = hex_to_rgb("888888")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[画像: {elem.get('alt', img_path)}]")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb("999999")


def _render_toc(doc, template: dict, toc_headings: list[dict]):
    toc_cfg = template.get("heading2", {})
    p = doc.add_paragraph()
    run = p.add_run("目次")
    run.bold = True
    run.font.name = toc_cfg.get("font", "Arial")
    run.font.size = Pt(14)
    run.font.color.rgb = hex_to_rgb(toc_cfg.get("color", "1B2A4A"))
    p.paragraph_format.space_after = Pt(8)

    body_cfg = template.get("body", {})
    for h in toc_headings:
        p = doc.add_paragraph()
        indent = {2: 0.5, 3: 1.5, 4: 2.5, 5: 3.5}.get(h["level"], 0.5)
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h["text"])
        run.font.name = body_cfg.get("font", "Arial")
        run.font.size = Pt(10)
        run.font.color.rgb = hex_to_rgb(body_cfg.get("color", "333333"))
        if h["level"] == 2:
            run.bold = True
